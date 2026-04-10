import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

import yaml


WORKSPACE_ROOT = Path(__file__).resolve().parents[5]
AUDIT_ROOT = WORKSPACE_ROOT / "construction-aduit"
S4_SCRIPT = AUDIT_ROOT / ".opencode/skills/construction-audit-s4-error-report/scripts/run_error_report.py"


def extract_docx_text(docx_path: Path) -> str:
    document = Document(docx_path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    return "\n".join(parts)


class RunErrorReportTests(unittest.TestCase):
    def test_fails_when_findings_directory_missing(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            config_path = tmpdir_path / "audit-config.yaml"
            config_payload = {
                "audit_id": "AUDIT-001",
                "audit_type": "budget",
                "project_info": {"project_name": "测试项目"},
                "audit_date": "2026-04-01",
                "output_dir": str(tmpdir_path),
            }
            config_path.write_text(yaml.safe_dump(config_payload, allow_unicode=True), encoding="utf-8")

            result = subprocess.run(
                [sys.executable, str(S4_SCRIPT), "--config", str(config_path)],
                capture_output=True,
                text=True,
                cwd=WORKSPACE_ROOT,
            )

            self.assertNotEqual(result.returncode, 0)
            self.assertIn("findings", result.stderr)

    def test_generates_markdown_and_docx_without_formal_json(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            findings_dir = tmpdir_path / "findings"
            findings_dir.mkdir()
            config_path = tmpdir_path / "audit-config.yaml"

            config_payload = {
                "audit_id": "AUDIT-001",
                "audit_type": "budget",
                "project_info": {"project_name": "测试项目"},
                "audit_date": "2026-04-01",
                "output_dir": str(tmpdir_path),
            }
            config_path.write_text(yaml.safe_dump(config_payload, allow_unicode=True), encoding="utf-8")

            findings_payload = {
                "sheet_name": "表一",
                "total_cells_checked": 3,
                "findings": [
                    {
                        "finding_id": "F-001",
                        "cell_ref": "C3",
                        "category": "费率",
                        "severity": "critical",
                        "actual_value": "2%",
                        "expected_value": "1.5%",
                        "discrepancy_pct": 0.5,
                        "description": "安全生产费率偏高",
                        "rule_source_anchor": "表一审核规则 > 安全生产费",
                        "rule_source_excerpt": "安全生产费 = 建筑安装工程费 * 1.5%",
                        "auto_fixable": True,
                    }
                ],
                "summary": {
                    "critical": 1,
                    "high": 0,
                    "medium": 0,
                    "low": 0,
                    "pass": 2,
                    "cumulative_deviation_pct": 2.5,
                    "cumulative_deviation_warning": False,
                },
            }
            (findings_dir / "findings_表一.json").write_text(
                json.dumps(findings_payload, ensure_ascii=False),
                encoding="utf-8",
            )

            result = subprocess.run(
                [sys.executable, str(S4_SCRIPT), "--config", str(config_path)],
                capture_output=True,
                text=True,
                cwd=WORKSPACE_ROOT,
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertIn("output_markdown=", result.stdout)
            self.assertIn("output_docx=", result.stdout)

            docx_path = tmpdir_path / "audit-report.docx"
            markdown_path = tmpdir_path / "audit-report.md"
            self.assertTrue(docx_path.exists())
            self.assertTrue(markdown_path.exists())
            self.assertTrue(markdown_path.read_text(encoding="utf-8").strip())
            self.assertFalse((tmpdir_path / "audit-report.json").exists())

            text = extract_docx_text(docx_path)
            self.assertIn("工程建设审核报告", text)
            self.assertIn("测试项目", text)
            self.assertIn("安全生产费率偏高", text)
            markdown = markdown_path.read_text(encoding="utf-8")
            self.assertIn("工程建设审核报告", markdown)
            self.assertIn("测试项目", markdown)
            self.assertIn("安全生产费率偏高", markdown)

    def test_aggregates_multiple_findings_files_into_one_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            findings_dir = tmpdir_path / "findings"
            findings_dir.mkdir()
            config_path = tmpdir_path / "audit-config.yaml"

            config_payload = {
                "audit_id": "AUDIT-002",
                "audit_type": "budget",
                "project_info": {"project_name": "多表测试项目"},
                "audit_date": "2026-04-01",
                "output_dir": str(tmpdir_path),
            }
            config_path.write_text(yaml.safe_dump(config_payload, allow_unicode=True), encoding="utf-8")

            findings_a = {
                "sheet_name": "表一",
                "total_cells_checked": 3,
                "findings": [
                    {
                        "finding_id": "F-001",
                        "cell_ref": "C3",
                        "category": "费率",
                        "severity": "critical",
                        "actual_value": "2%",
                        "expected_value": "1.5%",
                        "discrepancy_pct": 0.5,
                        "description": "安全生产费率偏高",
                        "rule_source_anchor": "表一审核规则 > 安全生产费",
                        "rule_source_excerpt": "安全生产费 = 建筑安装工程费 * 1.5%",
                        "auto_fixable": False,
                    }
                ],
                "summary": {
                    "critical": 1,
                    "high": 0,
                    "medium": 0,
                    "low": 0,
                    "pass": 2,
                    "cumulative_deviation_pct": 2.5,
                    "cumulative_deviation_warning": False,
                },
            }
            findings_b = {
                "sheet_name": "表二",
                "total_cells_checked": 2,
                "findings": [
                    {
                        "finding_id": "F-002",
                        "cell_ref": "D5",
                        "category": "合计",
                        "severity": "medium",
                        "actual_value": "100",
                        "expected_value": "120",
                        "discrepancy_pct": 20.0,
                        "description": "合计不符",
                        "rule_source_anchor": "表二审核规则 > 合计",
                        "rule_source_excerpt": "合计 = 左值 + 右值",
                        "auto_fixable": True,
                    }
                ],
                "summary": {
                    "critical": 0,
                    "high": 0,
                    "medium": 1,
                    "low": 0,
                    "pass": 1,
                    "cumulative_deviation_pct": 1.2,
                    "cumulative_deviation_warning": False,
                },
            }
            (findings_dir / "findings_表一.json").write_text(json.dumps(findings_a, ensure_ascii=False), encoding="utf-8")
            (findings_dir / "findings_表二.json").write_text(json.dumps(findings_b, ensure_ascii=False), encoding="utf-8")

            result = subprocess.run(
                [sys.executable, str(S4_SCRIPT), "--config", str(config_path)],
                capture_output=True,
                text=True,
                cwd=WORKSPACE_ROOT,
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            text = extract_docx_text(tmpdir_path / "audit-report.docx")
            self.assertIn("多表测试项目", text)
            self.assertIn("表一", text)
            self.assertIn("表二", text)
            self.assertIn("安全生产费率偏高", text)
            self.assertIn("合计不符", text)

    def test_reports_only_l7_l8_l9_for_sample_regression_gate(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            findings_dir = tmpdir_path / "findings"
            findings_dir.mkdir()
            config_path = tmpdir_path / "audit-config.yaml"

            config_payload = {
                "audit_id": "AUDIT-L789-001",
                "audit_type": "budget",
                "project_info": {"project_name": "东海县海陵家苑三网小区新建工程"},
                "audit_date": "2026-04-01",
                "output_dir": str(tmpdir_path),
            }
            config_path.write_text(yaml.safe_dump(config_payload, allow_unicode=True), encoding="utf-8")

            findings_payload = {
                "sheet_name": "表一（451定额折前）",
                "total_cells_checked": 9,
                "findings": [
                    {
                        "finding_id": "F-L7",
                        "cell_ref": "L7",
                        "category": "含税价",
                        "severity": "critical",
                        "actual_value": "11659.81",
                        "expected_value": "10964.5973412",
                        "description": "建筑安装工程费含税价不符",
                        "rule_source_anchor": "表一（451定额折前）审核规则/建筑安装工程费",
                        "rule_source_excerpt": "建筑安装工程费 | 安全生产费计算表（表二折前）建筑安装工程费所对应的合计（元）",
                        "auto_fixable": False,
                    },
                    {
                        "finding_id": "F-L8",
                        "cell_ref": "L8",
                        "category": "含税价",
                        "severity": "critical",
                        "actual_value": "91.64",
                        "expected_value": "0.0",
                        "description": "需要安装的设备费含税价不符",
                        "rule_source_anchor": "表一（451定额折前）审核规则/需要安装的设备费",
                        "rule_source_excerpt": "需要安装的设备费 | 表四-需安装设备所对应的合计（元）-除税价",
                        "auto_fixable": False,
                    },
                    {
                        "finding_id": "F-L9",
                        "cell_ref": "L9",
                        "category": "含税价",
                        "severity": "critical",
                        "actual_value": "3871.41",
                        "expected_value": "802.5848711742",
                        "description": "工程建设其他费含税价不符",
                        "rule_source_anchor": "表一（451定额折前）审核规则/工程建设其他费",
                        "rule_source_excerpt": "工程建设其他费 | 表五（甲）...总计",
                        "auto_fixable": False,
                    },
                ],
                "summary": {
                    "critical": 3,
                    "high": 0,
                    "medium": 0,
                    "low": 0,
                    "pass": 6,
                    "cumulative_deviation_pct": 0.0,
                    "cumulative_deviation_warning": False,
                },
            }
            (findings_dir / "findings_表一_451定额折前_.json").write_text(
                json.dumps(findings_payload, ensure_ascii=False),
                encoding="utf-8",
            )

            result = subprocess.run(
                [sys.executable, str(S4_SCRIPT), "--config", str(config_path)],
                capture_output=True,
                text=True,
                cwd=WORKSPACE_ROOT,
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertIn("total_findings=3", result.stdout)
            self.assertIn("output_markdown=", result.stdout)
            text = extract_docx_text(tmpdir_path / "audit-report.docx")
            self.assertIn("L7", text)
            self.assertIn("L8", text)
            self.assertIn("L9", text)
            markdown = (tmpdir_path / "audit-report.md").read_text(encoding="utf-8")
            self.assertIn("L7", markdown)
            self.assertIn("L8", markdown)
            self.assertIn("L9", markdown)


if __name__ == "__main__":
    unittest.main()
