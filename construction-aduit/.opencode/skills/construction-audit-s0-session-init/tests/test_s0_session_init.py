import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from shutil import rmtree, which
from uuid import uuid4

import openpyxl
import xlrd
import yaml
from docx import Document


WORKSPACE_ROOT = Path(__file__).resolve().parents[5]
AUDIT_ROOT = WORKSPACE_ROOT / "construction-aduit"
SCRIPT = AUDIT_ROOT / ".opencode/skills/construction-audit-s0-session-init/scripts/session_init.py"


class S0SessionInitCliTests(unittest.TestCase):
    def _make_rule_document(self, path: Path) -> None:
        document = Document()
        document.add_heading("审核规则", level=1)
        document.add_paragraph("人工费应等于技工费加普工费。")
        document.save(path)

    def _make_xlsx(self, path: Path) -> None:
        workbook = openpyxl.Workbook()
        workbook.active.title = "表一"
        workbook.active["A1"] = "预算"
        workbook.create_sheet("表二")
        workbook.save(path)

    def _make_xlsx_with_hidden_sheet(self, path: Path) -> None:
        workbook = openpyxl.Workbook()
        workbook.active.title = "工程信息"
        workbook.active["A1"] = "预算"
        workbook.create_sheet("隐藏费率表")
        workbook["隐藏费率表"].sheet_state = "hidden"
        workbook.create_sheet("表一")
        workbook.save(path)

    def _make_xls(self, path: Path) -> None:
        python = which("python") or sys.executable
        script = f"""
import xlwt

workbook = xlwt.Workbook()
sheet = workbook.add_sheet("表一")
sheet.write(0, 0, "结算")
workbook.add_sheet("表二")
workbook.save(r"{path}")
"""
        result = subprocess.run(
            [python, "-c", script],
            capture_output=True,
            text=True,
            cwd=WORKSPACE_ROOT,
        )
        if result.returncode != 0:
            raise AssertionError(result.stderr or result.stdout)

    def _run(
        self,
        rule_path: Path,
        spreadsheet_path: Path,
        output_dir: Path,
        *,
        sheet_scope: str | None = None,
    ):
        command = [
            sys.executable,
            str(SCRIPT),
            "--rule-document",
            str(rule_path),
            "--spreadsheet",
            str(spreadsheet_path),
            "--audit-type",
            "budget",
            "--project-name",
            "东海县海陵家苑三网小区新建工程",
            "--output-dir",
            str(output_dir),
        ]
        if sheet_scope:
            command.extend(["--sheet-scope", sheet_scope])
        return subprocess.run(
            command,
            capture_output=True,
            text=True,
            cwd=WORKSPACE_ROOT,
        )

    def _assert_config_complete(self, config: dict, expected_source_format: str) -> None:
        self.assertIn("audit_id", config)
        self.assertIn("audit_type", config)
        self.assertIn("audit_date", config)
        self.assertIn("project_info", config)
        self.assertIn("rule_document", config)
        self.assertIn("spreadsheet", config)
        self.assertIn("output_dir", config)
        self.assertIn("markdown_path", config["rule_document"])

        self.assertEqual(config["audit_type"], "budget")
        self.assertEqual(config["project_info"]["project_name"], "东海县海陵家苑三网小区新建工程")
        self.assertTrue(str(config["spreadsheet"]["working_path"]).endswith(f".{expected_source_format}"))
        self.assertEqual(config["spreadsheet"]["source_format"], expected_source_format)
        self.assertEqual(config["spreadsheet"]["path"], config["spreadsheet"]["working_path"])
        self.assertIn("sheets", config["spreadsheet"])
        self.assertIn("sheet_scope", config["spreadsheet"])
        self.assertIn("all_sheets", config["spreadsheet"])
        self.assertIn("visible_sheets", config["spreadsheet"])
        self.assertIn("hidden_sheets", config["spreadsheet"])
        self.assertTrue(config["spreadsheet"]["sheets"])
        self.assertTrue(str(config["rule_document"]["markdown_path"]).endswith("/rule_doc.md"))

    def test_normalizes_real_xlsx_without_changing_source_format(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            rule_path = tmpdir_path / "rule_document.docx"
            spreadsheet_path = tmpdir_path / "spreadsheet.xlsx"
            output_dir = tmpdir_path / "audit-output"
            config_path = output_dir / "audit-config.yaml"

            self._make_rule_document(rule_path)
            self._make_xlsx(spreadsheet_path)

            result = self._run(rule_path, spreadsheet_path, output_dir)

            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertIn("audit_id=", result.stdout)
            self.assertIn("source_format=xlsx", result.stdout)
            self.assertIn("target_sheets_count=2", result.stdout)

            config = yaml.safe_load(config_path.read_text(encoding="utf-8"))
            self._assert_config_complete(config, "xlsx")
            self.assertEqual(config["spreadsheet"]["original_path"], str(spreadsheet_path.resolve()))
            self.assertTrue(Path(config["spreadsheet"]["working_path"]).exists())
            self.assertEqual(
                openpyxl.load_workbook(config["spreadsheet"]["working_path"]).sheetnames,
                ["表一", "表二"],
            )

    def test_preserves_real_xls_as_working_xls(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            rule_path = tmpdir_path / "rule_document.docx"
            spreadsheet_path = tmpdir_path / "spreadsheet.xls"
            output_dir = tmpdir_path / "audit-output"
            config_path = output_dir / "audit-config.yaml"

            self._make_rule_document(rule_path)
            self._make_xls(spreadsheet_path)

            result = self._run(rule_path, spreadsheet_path, output_dir)

            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertIn("audit_id=", result.stdout)
            self.assertIn("source_format=xls", result.stdout)
            self.assertIn("target_sheets_count=2", result.stdout)

            config = yaml.safe_load(config_path.read_text(encoding="utf-8"))
            self._assert_config_complete(config, "xls")
            self.assertEqual(config["spreadsheet"]["original_path"], str(spreadsheet_path.resolve()))
            self.assertTrue(Path(config["spreadsheet"]["working_path"]).exists())
            workbook = xlrd.open_workbook(config["spreadsheet"]["working_path"])
            self.assertEqual(workbook.sheet_names(), ["表一", "表二"])

    def test_accepts_excel_content_with_wrong_extension(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            rule_path = tmpdir_path / "rule_document.docx"
            spreadsheet_path = tmpdir_path / "spreadsheet.bin"
            output_dir = tmpdir_path / "audit-output"
            config_path = output_dir / "audit-config.yaml"

            self._make_rule_document(rule_path)
            self._make_xlsx(tmpdir_path / "source.xlsx")
            (tmpdir_path / "source.xlsx").replace(spreadsheet_path)

            result = self._run(rule_path, spreadsheet_path, output_dir)

            self.assertEqual(result.returncode, 0, result.stderr)
            config = yaml.safe_load(config_path.read_text(encoding="utf-8"))
            self._assert_config_complete(config, "xlsx")
            self.assertEqual(config["spreadsheet"]["original_path"], str(spreadsheet_path.resolve()))
            self.assertTrue(Path(config["spreadsheet"]["working_path"]).exists())
            self.assertEqual(
                openpyxl.load_workbook(config["spreadsheet"]["working_path"]).sheetnames,
                ["表一", "表二"],
            )

    def test_rejects_non_excel_input_with_friendly_error(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            rule_path = tmpdir_path / "rule_document.docx"
            spreadsheet_path = tmpdir_path / "spreadsheet.txt"
            output_dir = tmpdir_path / "audit-output"

            self._make_rule_document(rule_path)
            spreadsheet_path.write_text("not an excel file", encoding="utf-8")

            result = self._run(rule_path, spreadsheet_path, output_dir)

            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Excel", result.stderr)
            self.assertFalse((output_dir / "audit-config.yaml").exists())

    def test_uses_visible_sheets_as_default_target_list(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            rule_path = tmpdir_path / "rule_document.docx"
            spreadsheet_path = tmpdir_path / "spreadsheet.xlsx"
            output_dir = tmpdir_path / "audit-output"
            config_path = output_dir / "audit-config.yaml"

            self._make_rule_document(rule_path)
            self._make_xlsx_with_hidden_sheet(spreadsheet_path)

            result = self._run(rule_path, spreadsheet_path, output_dir)

            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertIn("target_sheets_count=2", result.stdout)
            config = yaml.safe_load(config_path.read_text(encoding="utf-8"))
            self._assert_config_complete(config, "xlsx")
            self.assertEqual(config["spreadsheet"]["all_sheets"], ["工程信息", "隐藏费率表", "表一"])
            self.assertEqual(config["spreadsheet"]["visible_sheets"], ["工程信息", "表一"])
            self.assertEqual(config["spreadsheet"]["hidden_sheets"], ["隐藏费率表"])
            self.assertEqual(config["spreadsheet"]["sheets"], ["工程信息", "表一"])
            self.assertEqual(config["spreadsheet"]["sheet_scope"], "visible")

    def test_rejects_sheet_scope_all_to_keep_hidden_sheets_out_of_direct_targets(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            rule_path = tmpdir_path / "rule_document.docx"
            spreadsheet_path = tmpdir_path / "spreadsheet.xlsx"
            output_dir = tmpdir_path / "audit-output"

            self._make_rule_document(rule_path)
            self._make_xlsx_with_hidden_sheet(spreadsheet_path)

            result = self._run(rule_path, spreadsheet_path, output_dir, sheet_scope="all")

            self.assertNotEqual(result.returncode, 0)
            self.assertIn("sheet_scope=visible", result.stderr)
            self.assertFalse((output_dir / "audit-config.yaml").exists())

    def test_supports_literal_tmp_output_dir(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            rule_path = tmpdir_path / "rule_document.docx"
            spreadsheet_path = tmpdir_path / "spreadsheet.xlsx"
            output_dir = Path("/tmp") / f"construction-audit-s0-tmp-{uuid4().hex}"
            config_path = output_dir / "audit-config.yaml"

            self._make_rule_document(rule_path)
            self._make_xlsx(spreadsheet_path)

            try:
                result = self._run(rule_path, spreadsheet_path, output_dir)

                self.assertEqual(result.returncode, 0, result.stderr)
                self.assertTrue(config_path.exists())
                config = yaml.safe_load(config_path.read_text(encoding="utf-8"))
                self._assert_config_complete(config, "xlsx")
                self.assertEqual(config["spreadsheet"]["original_path"], str(spreadsheet_path.resolve()))
                self.assertIn("config_path=", result.stdout)
            finally:
                if output_dir.exists():
                    rmtree(output_dir)


if __name__ == "__main__":
    unittest.main()
