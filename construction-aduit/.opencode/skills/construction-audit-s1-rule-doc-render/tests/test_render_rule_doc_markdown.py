import subprocess
import tempfile
import unittest
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[5]
WORKSPACE_ROOT = REPO_ROOT
AUDIT_ROOT = REPO_ROOT / "construction-aduit"
SCRIPT = AUDIT_ROOT / ".opencode/skills/construction-audit-s1-rule-doc-render/scripts/render_rule_doc_markdown.sh"


class RenderRuleDocMarkdownTests(unittest.TestCase):
    def _run(self, input_path: Path, output_path: Path) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [str(SCRIPT), str(input_path), str(output_path)],
            capture_output=True,
            text=True,
            cwd=WORKSPACE_ROOT,
        )

    def test_renders_docx_to_markdown(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            input_path = tmpdir_path / "rule_document.docx"
            output_path = tmpdir_path / "rule_doc.md"

            document = Document()
            document.add_heading("表一审核规则", level=1)
            document.add_paragraph("预备费按三项费用之和的0.4%计取。")
            document.save(input_path)

            result = self._run(input_path, output_path)

            self.assertEqual(result.returncode, 0, result.stderr)
            content = output_path.read_text(encoding="utf-8")
            self.assertIn("表一审核规则", content)
            self.assertIn("预备费按三项费用之和的0.4%计取。", content)

    def test_preserves_table_as_markdown_table(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            input_path = tmpdir_path / "rule_document.docx"
            output_path = tmpdir_path / "rule_doc.md"

            document = Document()
            document.add_heading("表二审核规则", level=1)
            table = document.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "费用名称"
            table.cell(0, 1).text = "依据和计算方法"
            table.cell(1, 0).text = "预备费"
            table.cell(1, 1).text = "三项费用之和*0.4%"
            table.cell(2, 0).text = "总计"
            table.cell(2, 1).text = "合计+预备费"
            document.save(input_path)

            result = self._run(input_path, output_path)

            self.assertEqual(result.returncode, 0, result.stderr)
            content = output_path.read_text(encoding="utf-8")
            self.assertIn("|", content)
            self.assertIn("费用名称", content)
            self.assertIn("依据和计算方法", content)
            self.assertIn("预备费", content)
            self.assertIn("合计+预备费", content)

    def test_fails_when_input_docx_missing(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            input_path = tmpdir_path / "missing.docx"
            output_path = tmpdir_path / "rule_doc.md"

            result = self._run(input_path, output_path)

            self.assertNotEqual(result.returncode, 0)
            self.assertIn("input docx not found", result.stderr)

    def test_fails_when_input_is_not_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            input_path = tmpdir_path / "rule_document.txt"
            output_path = tmpdir_path / "rule_doc.md"
            input_path.write_text("not a docx", encoding="utf-8")

            result = self._run(input_path, output_path)

            self.assertNotEqual(result.returncode, 0)
            self.assertIn("input must be a .docx file", result.stderr)


if __name__ == "__main__":
    unittest.main()
