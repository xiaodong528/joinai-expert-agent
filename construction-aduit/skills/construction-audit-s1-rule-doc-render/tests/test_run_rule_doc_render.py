import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from shutil import rmtree
from uuid import uuid4

import yaml
from docx import Document
from openpyxl import Workbook


WORKSPACE_ROOT = Path(__file__).resolve().parents[5]
AUDIT_ROOT = WORKSPACE_ROOT / "joinai-expert-agent/construction-aduit"
CONSTRUCTION_REVIEW_ROOT = WORKSPACE_ROOT / "construction-review"
SCRIPT = AUDIT_ROOT / "skills/construction-audit-s1-rule-doc-render/scripts/run_rule_doc_render.py"
REAL_RULE_DOC = CONSTRUCTION_REVIEW_ROOT / "examples/rules-docx/家客预算审核知识库11.9.docx"


def extract_budget_fee_rows(section_text: str) -> list[str]:
    rows: list[str] = []
    for line in section_text.splitlines():
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        inner = stripped[1:-1] if stripped.endswith("|") else stripped[1:]
        cells = [cell.strip() for cell in inner.split("|")]
        if len(cells) < 2:
            continue
        if cells[0] in {"", "费用名称"}:
            continue
        if set(cells[0]) <= {"-", ":", " "}:
            continue
        rows.append(cells[0])
    return rows


class RunRuleDocRenderTests(unittest.TestCase):
    def _make_rule_document(self, path: Path) -> None:
        document = Document()
        document.add_heading("表一（451定额折前）审核规则", level=1)
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "费用名称"
        table.cell(0, 1).text = "依据和计算方法"
        table.cell(1, 0).text = "预备费"
        table.cell(1, 1).text = "三项费用之和*0.4%"
        table.cell(2, 0).text = "总计"
        table.cell(2, 1).text = "合计+预备费"
        document.save(path)

    def _write_config(self, path: Path, config: dict) -> None:
        path.write_text(yaml.safe_dump(config, allow_unicode=True, sort_keys=False), encoding="utf-8")

    def _run(self, config_path: Path) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(SCRIPT), "--config", str(config_path)],
            capture_output=True,
            text=True,
            cwd=WORKSPACE_ROOT,
        )

    def test_renders_from_minimal_audit_config(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            input_path = tmpdir_path / "rule_document.docx"
            output_path = tmpdir_path / "rule_doc.md"
            config_path = tmpdir_path / "audit-config.yaml"

            self._make_rule_document(input_path)
            self._write_config(
                config_path,
                {
                    "audit_id": "AUDIT-TEST-001",
                    "rule_document": {
                        "path": str(input_path),
                        "markdown_path": str(output_path),
                    },
                },
            )

            result = self._run(config_path)

            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertTrue(output_path.exists())
            self.assertIn("audit_id=AUDIT-TEST-001", result.stdout)
            self.assertIn(f"input_docx={input_path.resolve()}", result.stdout)
            self.assertIn(f"output_markdown={output_path.resolve()}", result.stdout)
            content = output_path.read_text(encoding="utf-8")
            self.assertIn("表一（451定额折前）审核规则", content)
            self.assertIn("预备费", content)
            self.assertIn("总计", content)
            self.assertLess(content.index("预备费"), content.index("总计"))

    def test_fails_when_rule_document_path_missing(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            config_path = tmpdir_path / "audit-config.yaml"
            self._write_config(
                config_path,
                {
                    "audit_id": "AUDIT-TEST-002",
                    "rule_document": {
                        "markdown_path": str(tmpdir_path / "rule_doc.md"),
                    },
                },
            )

            result = self._run(config_path)

            self.assertNotEqual(result.returncode, 0)
            self.assertIn("rule_document.path", result.stderr)

    def test_fails_when_markdown_path_missing(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            input_path = tmpdir_path / "rule_document.docx"
            config_path = tmpdir_path / "audit-config.yaml"

            self._make_rule_document(input_path)
            self._write_config(
                config_path,
                {
                    "audit_id": "AUDIT-TEST-003",
                    "rule_document": {
                        "path": str(input_path),
                    },
                },
            )

            result = self._run(config_path)

            self.assertNotEqual(result.returncode, 0)
            self.assertIn("rule_document.markdown_path", result.stderr)

    def test_supports_literal_tmp_config_path_from_s0_chain(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            input_path = tmpdir_path / "rule_document.docx"
            spreadsheet_path = tmpdir_path / "spreadsheet.xlsx"
            s0_script = AUDIT_ROOT / "skills/construction-audit-s0-session-init/scripts/session_init.py"
            output_dir = Path("/tmp") / f"construction-audit-s1-tmp-{uuid4().hex}"
            config_path = output_dir / "audit-config.yaml"

            self._make_rule_document(input_path)
            workbook_path = spreadsheet_path
            excel = Workbook()
            excel.active.title = "工程信息"
            excel.active["A1"] = "预算"
            excel.save(workbook_path)

            try:
                s0_result = subprocess.run(
                    [
                        sys.executable,
                        str(s0_script),
                        "--rule-document",
                        str(input_path),
                        "--spreadsheet",
                        str(workbook_path),
                        "--audit-type",
                        "budget",
                        "--project-name",
                        "东海县海陵家苑三网小区新建工程",
                        "--output-dir",
                        str(output_dir),
                    ],
                    capture_output=True,
                    text=True,
                    cwd=WORKSPACE_ROOT,
                )

                self.assertEqual(s0_result.returncode, 0, s0_result.stderr)
                self.assertTrue(config_path.exists())

                result = self._run(config_path)

                self.assertEqual(result.returncode, 0, result.stderr)
                self.assertIn("audit_id=", result.stdout)
                self.assertIn(f"input_docx={input_path.resolve()}", result.stdout)
                content = (output_dir / "rule_doc.md").read_text(encoding="utf-8")
                self.assertIn("表一（451定额折前）审核规则", content)
            finally:
                if output_dir.exists():
                    rmtree(output_dir)

    def test_real_rule_doc_keeps_nine_budget_fee_rows_in_order(self):
        if not REAL_RULE_DOC.exists():
            self.skipTest(f"Sample rule doc not found: {REAL_RULE_DOC}")

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            output_path = tmpdir_path / "rule_doc.md"
            config_path = tmpdir_path / "audit-config.yaml"
            self._write_config(
                config_path,
                {
                    "audit_id": "AUDIT-TEST-REAL-BUDGET-001",
                    "rule_document": {
                        "path": str(REAL_RULE_DOC),
                        "markdown_path": str(output_path),
                    },
                },
            )

            result = self._run(config_path)

            self.assertEqual(result.returncode, 0, result.stderr)
            content = output_path.read_text(encoding="utf-8")
            start = content.index("表一（451定额折前）审核规则")
            end = content.index("表一（综合工日折后）审核规则")
            budget_section = content[start:end]
            self.assertEqual(
                extract_budget_fee_rows(budget_section),
                [
                    "建筑安装工程费",
                    "需要安装的设备费",
                    "工程建设其他费",
                    "不需要安装的设备、工具费",
                    "小型建筑工程费",
                    "合计",
                    "预备费",
                    "总计",
                    "企业运营费",
                ],
            )


if __name__ == "__main__":
    unittest.main()
