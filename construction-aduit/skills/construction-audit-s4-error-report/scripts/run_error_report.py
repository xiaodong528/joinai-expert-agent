#!/usr/bin/env python3
"""
run_error_report.py - Stage 5 entrypoint for construction audit error report generation.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any

from report_inputs import load_findings, load_yaml
from report_payload import build_report_payload

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore


def friendly_error(message: str) -> int:
    print(f"Error: {message}", file=sys.stderr)
    return 1


def validate_stage_inputs(config_path: Path) -> tuple[dict[str, Any], Path, Path]:
    if not config_path.is_file():
        raise ValueError(f"配置文件不存在: {config_path}")
    config = load_yaml(str(config_path))
    if not isinstance(config, dict):
        raise ValueError(f"配置文件格式不正确: {config_path}")

    output_dir_raw = str(config.get("output_dir", "")).strip()
    if not output_dir_raw:
        raise ValueError("配置缺少 output_dir")
    output_dir = Path(output_dir_raw)

    findings_dir = output_dir / "findings"
    if not findings_dir.is_dir():
        raise ValueError(f"缺少 findings 目录: {findings_dir}")

    findings_files = sorted(findings_dir.glob("findings_*.json"))
    if not findings_files:
        raise ValueError(f"findings 目录为空: {findings_dir}")

    return config, output_dir, findings_dir


def build_markdown(report_payload: dict[str, Any], output_markdown: Path) -> None:
    project_info = report_payload.get("project_info", {})
    summary = report_payload.get("summary", {})

    lines: list[str] = [
        "# 工程建设审核报告",
        "",
        "## 一、基本信息",
        f"- 工程名称：{project_info.get('project_name', '')}",
        f"- 审核类型：{report_payload.get('audit_type', '')}",
        f"- 审核编号：{report_payload.get('audit_id', '')}",
        f"- 审核日期：{report_payload.get('audit_date', '')}",
        "",
        "## 二、审核摘要",
        f"- 总检查项数：{summary.get('total_checks', 0)}",
        f"- 发现问题数：{summary.get('total_findings', 0)}",
        f"- 通过项数：{summary.get('pass_count', 0)}",
        f"- 通过率：{summary.get('pass_rate', 0)}%",
    ]

    severity = summary.get("severity_totals", {})
    lines.append(
        "- 严重程度统计："
        f"critical={severity.get('critical', 0)}, "
        f"high={severity.get('high', 0)}, "
        f"medium={severity.get('medium', 0)}, "
        f"low={severity.get('low', 0)}"
    )
    lines.extend(["", "## 三、分表审核详情"])

    for sheet in report_payload.get("sheets", []):
        sheet_name = sheet.get("sheet_name", "未知表格")
        lines.extend(
            [
                "",
                f"### {sheet_name}",
                f"- 检查单元格数：{sheet.get('total_cells_checked', 0)}",
                f"- 发现问题数：{sheet.get('findings_count', 0)}",
                f"- 累计偏差率：{sheet.get('cumulative_deviation_pct', 0)}%",
            ]
        )
        findings = sheet.get("findings", [])
        if not findings:
            lines.append("- 本表格审核通过，未发现问题。")
            continue

        lines.extend(
            [
                "",
                "| 问题编号 | 单元格 | 类别 | 严重程度 | 实际值/期望值 | 说明 |",
                "|---|---|---|---|---|---|",
            ]
        )
        for finding in findings:
            lines.append(
                "| "
                + " | ".join(
                    [
                        str(finding.get("finding_id", "")),
                        str(finding.get("cell_ref", "")),
                        str(finding.get("category", "")),
                        str(finding.get("severity", "")),
                        f"{finding.get('actual_value', '')} / {finding.get('expected_value', '')}",
                        str(finding.get("description", "")),
                    ]
                )
                + " |"
            )

    lines.extend(["", "## 四、风险提示"])
    human_review_items = report_payload.get("human_review_items", [])
    if human_review_items:
        for item in human_review_items:
            lines.append(
                f"- {item.get('finding_id', '')} {item.get('sheet_name', '')} {item.get('cell_ref', '')} {item.get('description', '')}"
            )
    else:
        lines.append("- 无额外人工复核项。")

    output_markdown.parent.mkdir(parents=True, exist_ok=True)
    output_markdown.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def build_docx(report_payload: dict[str, Any], output_docx: Path) -> None:
    if Document is None:
        raise ValueError("缺少 python-docx，无法生成 audit-report.docx")

    document = Document()
    document.add_heading("工程建设审核报告", level=0)

    document.add_heading("一、基本信息", level=1)
    project_info = report_payload.get("project_info", {})
    for label, value in (
        ("工程名称", project_info.get("project_name", "")),
        ("审核类型", report_payload.get("audit_type", "")),
        ("审核编号", report_payload.get("audit_id", "")),
        ("审核日期", report_payload.get("audit_date", "")),
    ):
        document.add_paragraph(f"{label}：{value}")

    document.add_heading("二、审核摘要", level=1)
    summary = report_payload.get("summary", {})
    for label, value in (
        ("总检查项数", summary.get("total_checks", 0)),
        ("发现问题数", summary.get("total_findings", 0)),
        ("通过项数", summary.get("pass_count", 0)),
        ("通过率", f"{summary.get('pass_rate', 0)}%"),
    ):
        document.add_paragraph(f"{label}：{value}")

    severity = summary.get("severity_totals", {})
    document.add_paragraph(
        "严重程度统计："
        f"critical={severity.get('critical', 0)}, "
        f"high={severity.get('high', 0)}, "
        f"medium={severity.get('medium', 0)}, "
        f"low={severity.get('low', 0)}"
    )

    document.add_heading("三、分表审核详情", level=1)
    for sheet in report_payload.get("sheets", []):
        sheet_name = sheet.get("sheet_name", "未知表格")
        document.add_heading(str(sheet_name), level=2)
        document.add_paragraph(
            f"检查单元格数：{sheet.get('total_cells_checked', 0)}；"
            f"发现问题数：{sheet.get('findings_count', 0)}；"
            f"累计偏差率：{sheet.get('cumulative_deviation_pct', 0)}%"
        )

        findings = sheet.get("findings", [])
        if not findings:
            document.add_paragraph("本表格审核通过，未发现问题。")
            continue

        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["问题编号", "单元格", "类别", "严重程度", "实际值/期望值", "说明"]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header

        for finding in findings:
            row = table.add_row().cells
            row[0].text = str(finding.get("finding_id", ""))
            row[1].text = str(finding.get("cell_ref", ""))
            row[2].text = str(finding.get("category", ""))
            row[3].text = str(finding.get("severity", ""))
            row[4].text = f"{finding.get('actual_value', '')} / {finding.get('expected_value', '')}"
            row[5].text = str(finding.get("description", ""))

    document.add_heading("四、风险提示", level=1)
    human_review_items = report_payload.get("human_review_items", [])
    if human_review_items:
        for item in human_review_items:
            document.add_paragraph(
                f"{item.get('finding_id', '')} {item.get('sheet_name', '')} {item.get('cell_ref', '')} {item.get('description', '')}"
            )
    else:
        document.add_paragraph("无额外人工复核项。")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run stage 5 error report generation")
    parser.add_argument("--config", required=True, help="Path to audit-config.yaml")
    args = parser.parse_args(argv)

    try:
        config, output_dir, findings_dir = validate_stage_inputs(Path(args.config))
        sheet_results = load_findings(str(findings_dir))
        report_payload = build_report_payload(config, sheet_results)

        output_markdown = output_dir / "audit-report.md"
        output_docx = output_dir / "audit-report.docx"
        build_markdown(report_payload, output_markdown)
        build_docx(report_payload, output_docx)

        if not output_markdown.is_file() or output_markdown.stat().st_size == 0:
            return friendly_error(f"Markdown 输出为空: {output_markdown}")
        if not output_docx.is_file() or output_docx.stat().st_size == 0:
            return friendly_error(f"DOCX 输出为空: {output_docx}")

        print(
            f"audit_id={config.get('audit_id', '')} "
            f"findings_files={len(sheet_results)} "
            f"total_findings={report_payload['summary']['total_findings']} "
            f"output_markdown={output_markdown.resolve()} "
            f"output_docx={output_docx.resolve()}"
        )
        return 0
    except (ValueError, FileNotFoundError) as exc:
        return friendly_error(str(exc))


if __name__ == "__main__":
    raise SystemExit(main())
