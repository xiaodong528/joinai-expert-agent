"""
文档解析器 — 将 docx/md/pdf 剧本文档提取为纯文本

支持格式:
  - .md / .txt: 直接读取
  - .docx: python-docx 提取段落文本
  - .pdf: pypdf 提取文本

用法:
    python doc_parser.py --input 剧本.docx
    python doc_parser.py --input 剧本.docx --output 提取结果.md
"""

import argparse
import os
import sys


def extract_markdown(path: str) -> str:
    """读取 Markdown 或纯文本文件。"""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def extract_docx(path: str) -> str:
    """从 docx 提取段落文本。"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要安装 python-docx (pip install python-docx)", file=sys.stderr)
        sys.exit(1)

    doc = Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 保留标题层级
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                paragraphs.append(f"{'#' * level} {text}")
            else:
                paragraphs.append(text)

    # 提取表格内容
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paragraphs.append("\n".join(rows))

    return "\n\n".join(paragraphs)


def extract_pdf(path: str) -> str:
    """从 PDF 提取文本。"""
    try:
        from pypdf import PdfReader
    except ImportError:
        print("错误: 需要安装 pypdf (pip install pypdf)", file=sys.stderr)
        sys.exit(1)

    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"--- 第 {i + 1} 页 ---\n{text.strip()}")
    return "\n\n".join(pages)


def extract_text(path: str) -> str:
    """根据文件扩展名选择提取方法。"""
    ext = os.path.splitext(path)[1].lower()

    if ext in (".md", ".txt", ".yaml", ".yml"):
        return extract_markdown(path)
    elif ext == ".docx":
        return extract_docx(path)
    elif ext == ".pdf":
        return extract_pdf(path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}（支持 .md/.txt/.docx/.pdf）")


def main():
    parser = argparse.ArgumentParser(description="文档解析器：提取剧本/创意文档为纯文本")
    parser.add_argument("--input", required=True, help="输入文档路径（.docx/.md/.pdf）")
    parser.add_argument("--output", help="输出文本路径（默认打印到 stdout）")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    text = extract_text(args.input)

    if args.output:
        os.makedirs(os.path.dirname(args.output) or ".", exist_ok=True)
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"已提取 {len(text)} 字符 → {args.output}")
    else:
        print(text)


if __name__ == "__main__":
    main()
